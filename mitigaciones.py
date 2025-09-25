# mitigaciones.py
import os
from docx import Document

# Carpeta base del proyecto
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# Carpeta donde guardaste los .docx de mitigaciones
DOCS_DIR = os.path.join(BASE_DIR, "mitigaciones")

# Mapa tipo de ataque -> nombre de archivo .docx (dentro de ./mitigaciones)
_DOCX_BY_ATTACK = {
    "dos": "dos.docx",
    "ddos": "ddos.docx",
    "fuerza_bruta": "bruteforce.docx",     # fuerza bruta
    "login_sospechoso": "login.docx",
    "phishing": "phishing.docx",
}

_cache: dict[str, str] = {}

def _docx_to_text(path: str) -> str:
    """Convierte un DOCX a texto plano con saltos de línea (incluye párrafos y tablas)."""
    doc = Document(path)
    parts = []

    # Párrafos
    for p in doc.paragraphs:
        if p.text:
            parts.append(p.text.strip())

    # Tablas 
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append("  • " + " | ".join(cells))

    # Normaliza saltos
    text = "\n".join(parts)
    # Limpieza básica de líneas vacías múltiples
    lines = [ln.strip() for ln in text.splitlines()]
    clean = []
    prev_blank = False
    for ln in lines:
        if ln == "":
            if not prev_blank:
                clean.append("")
            prev_blank = True
        else:
            clean.append(ln)
            prev_blank = False
    return "\n".join(clean).strip()

def get_mitigaciones(attack_code: str) -> str:
    """
    Devuelve el texto de mitigaciones según el tipo de ataque.
    attack_code ∈ {'dos','ddos','fuerza_bruta','login_sospechoso','phishing'}
    """
    code = (attack_code or "").strip().lower()
    filename = _DOCX_BY_ATTACK.get(code)
    if not filename:
        return "Mitigaciones no disponibles para este tipo de ataque."

    if code in _cache:
        return _cache[code]

    path = os.path.join(DOCS_DIR, filename)
    if not os.path.exists(path):
        # Mensaje claro indicando la ruta esperada
        return f"Documento de mitigaciones no encontrado: {path}"

    try:
        text = _docx_to_text(path)
        _cache[code] = text if text else "Documento de mitigaciones vacío."
        return _cache[code]
    except Exception as e:
        return f"Error leyendo mitigaciones desde {path}: {e}"
